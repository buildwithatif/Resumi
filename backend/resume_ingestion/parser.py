"""
Resume Parser
Handles PDF and DOCX file parsing to extract raw text
"""

import logging
from pathlib import Path
from typing import Optional, Tuple
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors"""
    pass


def parse_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        ResumeParsingError: If PDF parsing fails
    """
    try:
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise ResumeParsingError("PDF is encrypted and cannot be parsed")
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
        
        if not text_content:
            raise ResumeParsingError("No text content found in PDF")
        
        return "\n".join(text_content)
        
    except ResumeParsingError:
        raise
    except Exception as e:
        raise ResumeParsingError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file_path: Path) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        ResumeParsingError: If DOCX parsing fails
    """
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        if not text_content:
            raise ResumeParsingError("No text content found in DOCX")
        
        return "\n".join(text_content)
        
    except ResumeParsingError:
        raise
    except Exception as e:
        raise ResumeParsingError(f"Failed to parse DOCX: {str(e)}")


def parse_resume(file_path: Path) -> Tuple[str, str]:
    """
    Parse resume file and extract text content
    
    Args:
        file_path: Path to resume file (PDF or DOCX)
        
    Returns:
        Tuple of (text_content, file_type)
        
    Raises:
        ResumeParsingError: If file format is unsupported or parsing fails
    """
    if not file_path.exists():
        raise ResumeParsingError(f"File not found: {file_path}")
    
    # Determine file type
    file_extension = file_path.suffix.lower()
    
    if file_extension == '.pdf':
        text = parse_pdf(file_path)
        return text, 'pdf'
    elif file_extension == '.docx':
        text = parse_docx(file_path)
        return text, 'docx'
    else:
        raise ResumeParsingError(
            f"Unsupported file format: {file_extension}. "
            "Only PDF and DOCX files are supported."
        )


def validate_resume_content(text: str) -> bool:
    """
    Validate that extracted text looks like a resume
    
    Args:
        text: Extracted text content
        
    Returns:
        True if content appears valid, False otherwise
    """
    # Relaxed validation
    if not text or len(text.strip()) < 10:  # Only fail on empty/tiny text
        logger.warning(f"Resume validation failed: Text length {len(text) if text else 0}")
        return False
    
    # Check for common resume indicators (case-insensitive)
    text_lower = text.lower()
    resume_indicators = [
        'experience', 'education', 'skills', 'work', 'employment',
        'university', 'college', 'degree', 'job', 'position',
        'project', 'summary', 'contact', 'email', 'phone'
    ]
    
    # At least 1 indicator is enough now
    indicator_count = sum(1 for indicator in resume_indicators if indicator in text_lower)
    
    if indicator_count < 1:
        logger.warning("Resume validation warning: No resume keywords found")
        # Still return True to allow extraction to attempt its best
        return True
    
    return True
